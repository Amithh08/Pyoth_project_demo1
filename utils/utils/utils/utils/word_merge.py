import os
from docx import Document

# Create output folder
OUTPUT_FOLDER = "outputs"
os.makedirs(OUTPUT_FOLDER, exist_ok=True)


def merge_word_documents(uploaded_files, output_filename="merged_document.docx"):
    """
    Merge multiple Word (.docx) documents into one.

    Parameters:
        uploaded_files (list): List of uploaded .docx files.
        output_filename (str): Output file name.

    Returns:
        str: Path to merged Word document.
    """

    if not uploaded_files:
        raise ValueError("No Word documents were uploaded.")

    merged_document = Document()

    # Remove the default empty paragraph
    if len(merged_document.paragraphs) > 0:
        p = merged_document.paragraphs[0]._element
        p.getparent().remove(p)

    for index, uploaded_file in enumerate(uploaded_files):

        uploaded_file.seek(0)
        document = Document(uploaded_file)

        # Add a heading before each document (optional)
        merged_document.add_heading(
            f"Document {index + 1}",
            level=2
        )

        # Copy all paragraphs
        for paragraph in document.paragraphs:
            merged_document.add_paragraph(paragraph.text)

        # Add a page break between documents
        if index != len(uploaded_files) - 1:
            merged_document.add_page_break()

    output_path = os.path.join(OUTPUT_FOLDER, output_filename)

    merged_document.save(output_path)

    return output_path
