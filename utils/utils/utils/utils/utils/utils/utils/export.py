import os
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document

# Output folder
OUTPUT_FOLDER = "outputs"

# Create folder if it doesn't exist
os.makedirs(OUTPUT_FOLDER, exist_ok=True)


def export_text_to_txt(text, filename="translated_text.txt"):
    """
    Export text to a TXT file.
    """

    filepath = os.path.join(OUTPUT_FOLDER, filename)

    with open(filepath, "w", encoding="utf-8") as file:
        file.write(text)

    return filepath


def export_text_to_docx(text, filename="translated_document.docx"):
    """
    Export text to a Word document.
    """

    filepath = os.path.join(OUTPUT_FOLDER, filename)

    document = Document()

    document.add_heading("Translated Document", level=1)

    document.add_paragraph(text)

    document.save(filepath)

    return filepath


def export_text_to_pdf(text, filename="translated_document.pdf"):
    """
    Export text to a PDF.
    """

    filepath = os.path.join(OUTPUT_FOLDER, filename)

    pdf = canvas.Canvas(filepath, pagesize=letter)

    width, height = letter

    x = 40
    y = height - 40

    for line in text.split("\n"):

        pdf.drawString(x, y, line)

        y -= 18

        # New page if needed
        if y < 40:
            pdf.showPage()
            y = height - 40

    pdf.save()

    return filepath


def export_existing_pdf(pdf_path):
    """
    Return an existing PDF path.
    """

    if not os.path.exists(pdf_path):
        raise FileNotFoundError("PDF file not found.")

    return pdf_path


def export_existing_audio(audio_path):
    """
    Return an existing MP3 path.
    """

    if not os.path.exists(audio_path):
        raise FileNotFoundError("Audio file not found.")

    return audio_path
